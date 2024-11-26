import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt

# from pygoogletranslation import Translator

# def translate(query):
#     proxy={"http":"127.0.0.1:7890","https":"127.0.0.1:7890"}
#     translator=Translator(proxies=proxy)
#     t=translator.translate(query ,src = 'en', dest = 'ru')
#     return t.text

# print(translate('hello').text)

# import pyGoogleTranslate

# pyGoogleTranslate.browser('chrome')
# pyGoogleTranslate.translate('Hello', 'ja')

# from googletrans import Translator
# translator = Translator(service_urls=[
#       'translate.google.cn', 'translate.google.com'])# 如果可以上外网，还可添加 'translate.google.com' 等
# trans=translator.translate('Hello World', src='en', dest='zh-cn')
# # 原文
# print(trans.origin)
# # 译文
# print(trans.text)

from translatepy.translators.google import GoogleTranslate

def translate(query):
    gtranslate = GoogleTranslate()
    res = gtranslate.translate(query, "ru")
    print(res)
    return str(res)

def json_to_docx(n0, json_file_path, docx_file_path):
    # 读取JSON文件内容
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    # 递归处理JSON数据并添加到DOCX文档中
    def process_data(n0, data):
        try:
            for k0 in data.keys():
                title_para = doc.add_heading(f'{n0+1}. '+translate(k0), level=1)
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for n1, k1 in enumerate(data[k0].keys()):
                    # 二级标题样式
                    title_para = doc.add_heading(f'{n0+1}.{n1+1} '+translate(k1), level=2)
                    for n2, k2 in enumerate(data[k0][k1]['content'].keys()):
                        # 三级标题样式
                        title_para = doc.add_heading(f'{n0+1}.{n1+1}.{n2+1} '+translate(k2), level=3)
                        para = doc.add_paragraph(translate(data[k0][k1]['content'][k2]['description']))
                        for n3, k3 in enumerate(data[k0][k1]['content'][k2]['content'].keys()):
                            # 四级标题样式
                            title_para = doc.add_heading(f'{n0+1}.{n1+1}.{n2+1}.{n3+1} '+translate(k3), level=4)
                            para = doc.add_paragraph(translate(data[k0][k1]['content'][k2]['content'][k3]['content']['svp']))
                            para = doc.add_paragraph(translate(data[k0][k1]['content'][k2]['content'][k3]['content']['parent_description']))
                            for item in data[k0][k1]['content'][k2]['content'][k3]['content']['description']:
                                para = doc.add_paragraph(translate(item), style='List Bullet')
        except Exception as e:
            print(e)


    process_data(n0, data)
    doc.save(docx_file_path)

def split_json_to_files(input_json_path):
    # 读取原始JSON文件内容
    with open(input_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    num = len(data.keys())
    for n0, (key, value) in enumerate(data.items()):
        # 构建每个单独JSON文件的文件名，这里简单以键名作为文件名，你可以按需修改
        file_name = f"temp{n0}.json"
        with open(file_name, 'w', encoding='utf-8') as out_f:
            # 将单个键值对重新构造成字典形式，以便保存为合法的JSON格式
            single_data = {key: value}
            json.dump(single_data, out_f, ensure_ascii=False, indent=4)
    return num

json_file_path = "all_content.json"  # 替换为实际的JSON文件路径
docx_file_path = "result_ru.docx"  # 替换为想要输出的DOCX文件路径
num = split_json_to_files(json_file_path)
for i in range(1, num):
    subjson_file_path = f"temp{i}.json"
    json_to_docx(i, subjson_file_path, f"temp{i}.docx")

from docxcompose.composer import Composer
from docx import Document
import os

def merge_docx_files(num, output_file):
    """
    合并多个DOCX文件到一个指定的输出文件中。

    :param input_files: 输入的DOCX文件列表，包含要合并的各个文件的路径。
    :param output_file: 合并后输出的DOCX文件的路径。
    """
    first_doc = Document("temp0.docx")
    composer = Composer(first_doc)

    for i in range(1, num):
        file_path = f"temp{i}.docx"
        doc_to_add = Document(file_path)
        composer.append(doc_to_add)

    composer.save(output_file)

output_file = "merged.docx"
merge_docx_files(num, output_file)
